# ---imports---
import docx
from docx2pdf import convert
import os
import pandas as pd
import threading
import queue
import tkinter as tk
from tkinter import filedialog as fd
from tkinter import ttk
# ---imports---

# ---variables---
num = 0
p_numbers = 0
placeholders = []
p_title = []
p_paragraph = []
save_with = ""
save_name = ""
save_folder = "."
doc_path = ""
data_path = ""
progress = 0
q = queue.Queue()
# ---variables---

# ---new window function---
def new():
    # ---close window function---
    def close():
        root.quit()
        return None
    # ---close window function---
    # ---check function---
    def check_thread():
        if thread.is_alive():
            if q.empty() == False:
                progress = q.get()
                bar["value"] = progress
                show.set(f"轉換完成檔案數: {progress} / {num}")
            window.after(100, check_thread)
        else:
            bar["value"] = num
            show.set(f"轉換完成檔案數: {num} / {num}\n已轉換完成。")
            wb_1 = tk.Button(window, text="關閉所有視窗", command = close)
            wb_1.pack()
    # ---check function---
    global progress, thread
    thread = threading.Thread(target=export)
    window = tk.Toplevel()
    window.title("輸出進度")
    window.minsize(width=200, height=100)
    window.resizable(False, False)
    bar = ttk.Progressbar(window, maximum=num)
    bar.pack()
    show = tk.StringVar()
    show.set(f"轉換完成檔案數: 0 / {num}")
    wp_1 = tk.Label(window, textvariable=show, font=("Microsoft JhengHei UI", 10))
    wp_1.pack()

    thread.start()
    check_thread()
# ---new window function---

# ---generating files function---
def export():
    data = pd.read_excel(data_path)
    for i in range(num):
        doc = docx.Document(doc_path)
        with q.mutex:
            q.queue.clear()
        q.put(i)
        for j in range(p_numbers):
            inline = doc.paragraphs[p_paragraph[j]].runs
            for k in range(len(inline)):
                if placeholders[j] in inline[k].text:
                    text = inline[k].text.replace(placeholders[j], data[p_title[j]][i])
                    inline[k].text = text
        docx_name = save_folder + "/" + save_name + "_" + data[save_with][i].replace(" ", "_") + ".docx"
        doc.save(docx_name)
        if typeVar.get():
            pdf_name = save_folder + "/"  + save_name + "_" + data[save_with][i].replace(" ", "_") + ".pdf"
            
            convert(docx_name, pdf_name)
            os.remove(docx_name)
    return None
# ---generating files function---

# ---placeholder function---
def find_words():
    doc = docx.Document(doc_path)
    p_location = [0] * p_numbers
    for i in range(len(doc.paragraphs)):
        for j in range(p_numbers):
            if doc.paragraphs[i].text.find(placeholders[j]) != -1:
                p_location[j] = i
    return p_location
# ---placeholder function---

# ---check placeholder styling function---
def check_run(index):
    doc = docx.Document(doc_path)
    inline = doc.paragraphs[p_paragraph[index]].runs
    for j in range(len(inline)):
        if placeholders[index] in inline[j].text:
            return True
    return False
# ---check placeholder styling function---

# ---choose doc function---
def select_doc():
    global doc_path
    filetypes = (('document files', '*.docx'),)
    doc_path = fd.askopenfilename(title='Choose a file', initialdir='./', filetypes=filetypes)
    show_doc.set(doc_path.split("/")[-1])
    output.set("")
    return None
# ---choose doc function---
    
# ---choose data function---
def select_data():
    global data_path
    filetypes = (('Excel files', '*.xlsx'),)
    data_path = fd.askopenfilename(title='Choose a file', initialdir='./', filetypes=filetypes)
    show_data.set(data_path.split("/")[-1])
    output.set("")
    return None
# ---choose data function---
    
# ---set word function---
def set_words():
    global p_numbers, placeholders, p_title
    if (i3_2.get() == "" or i3_2.get() == "" or i3_3.get() == ""):
        err = "error: 3. 請輸入關鍵字。"
        output.set(err)
        return None
    p_numbers = int(i3_1.get())
    placeholders = i3_2.get().split(",")
    p_title = i3_3.get().split(",")
    if (p_numbers == 0):
        err = "error: 3. 替代數量須為正整數。"
        output.set(err)
        return None
    elif (p_numbers == len(placeholders) and p_numbers == len(p_title)):
        output.set("")
        return None
    else:
        err = f"error: 3. 替代數量:{p_numbers}，替代文字數量:{len(placeholders)}，標題名稱數量:{len(p_title)}。三者不相符。"
        output.set(err)
        return None
# ---set word function---
    
# ---set save function---
def set_save():
    global save_with, save_name
    save_name = i4_1.get()
    save_with = i4_2.get()
    if save_name == "" or save_with == "":
        err = "error: 4. 請設定輸出檔案。"
        output.set(err)
        return None
    elif save_with in p_title:
        output.set("")
        return None
    else:
        err = "error: 4. 命名標題不存在於標題名稱中。"
        output.set(err)
        return None
# ---set save function---

# ---set save folder function---
def select_folder():
    global save_folder
    save_folder = fd.askdirectory(title="select a directory", initialdir='./')
    show_folder.set(save_folder)
# ---set save folder function---
    
# ---execute function---
def execute():
    global p_paragraph, num
    output.set("")
    #validity check
    if os.path.isfile(doc_path) == False:
        err = "error: 1. 範本檔案(.docx)不存在，請重新選擇。"
        output.set(err)
        return None
    if os.path.isfile(data_path) == False:
        err = "error: 2. 資料檔案(.xlsx)不存在，請重新選擇。"
        output.set(err)
        return None
    if (p_numbers == 0 or len(placeholders) == 0 or len(p_title) == 0):
        err = "error: 3. 請輸入關鍵字。"
        output.set(err)
        return None
    if (p_numbers == len(placeholders) and p_numbers == len(p_title)) == False:
        err = f"error: 3. 替代數量:{p_numbers}，替代文字數量:{len(placeholders)}，標題名稱數量:{len(p_title)}。三者不相符。"
        output.set(err)
        return None
    if save_name == "" or save_with == "":
        err = "error: 4. 請設定輸出檔案。"
        output.set(err)
        return None
    if save_with not in p_title:
        err = "error: 4. 命名標題不存在於標題名稱中。"
        output.set(err)
        return None
    if os.path.exists(save_folder) == False:
        err = "error: 4. 請選擇儲存資料夾。"
    #end of validity check

    p_paragraph = find_words()
    #existence check
    for i in range(len(p_paragraph)):
        if p_paragraph[i] == 0:
            err = "error: 範例檔案中該關鍵字不存在: " + placeholders[i]
            output.set(err)
            return None
    for i in range(len(p_paragraph)):
        if check_run(i) == False:
            err = "error: 範例檔案中該關鍵字格式不統一: " + placeholders[i]
            output.set(err)
            return None
    data = pd.read_excel(data_path)
    for i in range(len(p_title)):
        if p_title[i] not in data:
            err = "error: 資料檔案中該標題名稱不存在: " + p_title[i]
            output.set(err)
            return None
    #end of existence check
        
    num = len(data[p_title[0]])
    b6["state"] = "disabled"
    new()
    return None
# ---execute function---

# ---main function---
root = tk.Tk()
root.title("通知／證明產生器")
root.minsize(width=250, height=580)
root.resizable(False, False)

p1 = tk.Label(root, text="1. 選取範本檔案(.docx):", font=("Microsoft JhengHei UI", 12))
p1.grid(column=0, row=0, columnspan=2, sticky=tk.W)
b1 = tk.Button(root, text="選取檔案", command=select_doc)
b1.grid(column=0, row=1)
show_doc = tk.StringVar()
p1_1 = tk.Label(root, textvariable=show_doc, font=("Microsoft JhengHei UI", 10))
p1_1.grid(column=1, row=1, sticky=tk.W)
pm_1 = tk.Label(root, text="", font=(10))
pm_1.grid(column=0, row=2)

p2 = tk.Label(root, text="2. 選取資料檔案(.xlsx):", font=("Microsoft JhengHei UI", 12))
p2.grid(column=0, row=3, columnspan=2, sticky=tk.W)
b2 = tk.Button(root, text="選取檔案", command=select_data)
b2.grid(column=0, row=4)
show_data = tk.StringVar()
p1_1 = tk.Label(root, textvariable=show_data, font=("Microsoft JhengHei UI", 10))
p1_1.grid(column=1, row=4, sticky=tk.W)
pm_2 = tk.Label(root, text="", font=(10))
pm_2.grid(column=0, row=5)

p3 = tk.Label(root, text="3. 設定關鍵字:", font=("Microsoft JhengHei UI", 12))
p3.grid(column=0, row=6, columnspan=2, sticky=tk.W)
p3_1 = tk.Label(root, text="替代數量:", font=("Microsoft JhengHei UI", 10))
p3_1.grid(column=0, row=7)
i3_1 = tk.Entry(root)
i3_1.grid(column=1, row=7, sticky=tk.W)
p3_2 = tk.Label(root, text="替代文字:", font=("Microsoft JhengHei UI", 10))
p3_2.grid(column=0, row=8)
i3_2 = tk.Entry(root)
i3_2.grid(column=1, row=8, sticky=tk.W)
p3_3 = tk.Label(root, text="標題名稱:", font=("Microsoft JhengHei UI", 10))
p3_3.grid(column=0, row=9)
i3_3 = tk.Entry(root)
i3_3.grid(column=1, row=9, sticky=tk.W)
b3 = tk.Button(root, text="儲存", command=set_words)
b3.grid(column=0, row=10, ipadx=4)
pm_3 = tk.Label(root, text="", font=(10))
pm_3.grid(column=0, row=11)

p4 = tk.Label(root, text="4. 設定輸出檔案命名:", font=("Microsoft JhengHei UI", 12))
p4.grid(column=0, row=12, columnspan=2, sticky=tk.W)
p4_1 = tk.Label(root, text="檔案名稱:", font=("Microsoft JhengHei UI", 10))
p4_1.grid(column=0, row=13)
i4_1 = tk.Entry(root)
i4_1.grid(column=1, row=13, sticky=tk.W)
p4_2 = tk.Label(root, text="命名標題:", font=("Microsoft JhengHei UI", 10))
p4_2.grid(column=0, row=14)
i4_2 = tk.Entry(root)
i4_2.grid(column=1, row=14, sticky=tk.W)
b4_1 = tk.Button(root, text="選取儲存資料夾", command=select_folder)
b4_1.grid(column=0, row=15)
show_folder = tk.StringVar()
p4_3 = tk.Label(root, textvariable=show_folder, font=("Microsoft JhengHei UI", 10))
p4_3.grid(column=1, row=15, sticky=tk.W)
b4_2 = tk.Button(root, text="儲存", command=set_save)
b4_2.grid(column=0, row=16, ipadx=4, pady=2)
pm_4 = tk.Label(root, text="", font=(10))
pm_4.grid(column=0, row=17)

p5 = tk.Label(root, text="5. 選擇輸出檔案類型:", font=("Microsoft JhengHei UI", 12))
p5.grid(column=0, row=18, columnspan=2, sticky=tk.W)
typeVar = tk.BooleanVar()
radio1 = tk.Radiobutton(root, text=".docx檔", variable=typeVar, value=False)
radio2 = tk.Radiobutton(root, text=".pdf檔", variable=typeVar, value=True)
radio1.grid(column=0, row=19)
radio2.grid(column=1, row=19, sticky=tk.W)
pm_5 = tk.Label(root, text="", font=(10))
pm_5.grid(column=0, row=20)

b6 = tk.Button(root, text="開始執行", command=execute)
b6.grid(column=0, row=21, columnspan=2)
output = tk.StringVar()
p6 = tk.Label(root, textvariable=output, font=("Microsoft JhengHei UI", 10), fg='red')
p6.grid(column=0, row=22, columnspan=2, sticky=tk.W)

root.mainloop()
# ---main function---